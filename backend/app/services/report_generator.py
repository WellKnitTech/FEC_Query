import io
import os
import csv
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
from io import BytesIO, StringIO
import base64
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfgen import canvas

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import matplotlib
matplotlib.use('Agg')  # Non-interactive backend
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.figure import Figure
import numpy as np

from app.services.fec_client import FECClient
from app.services.analysis import AnalysisService
from app.services.fraud_detection import FraudDetectionService
from app.models.schemas import (
    CandidateSummary, FinancialSummary, ContributionAnalysis,
    ExpenditureBreakdown, EmployerAnalysis, ContributionVelocity,
    FraudAnalysis, MoneyFlowGraph
)


class ReportGenerator:
    """Generate candidate reports in PDF, DOCX, and Markdown formats"""
    
    def __init__(self, fec_client: FECClient):
        self.fec_client = fec_client
        self.analysis_service = AnalysisService(fec_client)
        self.fraud_service = FraudDetectionService(fec_client)
        
    async def collect_candidate_data(
        self,
        candidate_id: str,
        cycle: Optional[int] = None
    ) -> Dict[str, Any]:
        """Collect all data for a single candidate"""
        data = {
            'candidate': None,
            'financials': [],
            'contribution_analysis': None,
            'expenditure_breakdown': None,
            'employer_analysis': None,
            'velocity': None,
            'fraud_analysis': None,
            'money_flow': None,
        }
        
        try:
            # Get candidate info
            candidate = await self.fec_client.get_candidate(candidate_id)
            if candidate:
                data['candidate'] = CandidateSummary(
                    candidate_id=candidate.get('candidate_id', candidate_id),
                    name=candidate.get('name', candidate.get('candidate_name', 'Unknown')),
                    office=candidate.get('office'),
                    party=candidate.get('party'),
                    state=candidate.get('state'),
                    district=candidate.get('district'),
                    election_years=candidate.get('election_years'),
                    active_through=candidate.get('active_through')
                )
            
            # Get financials
            totals = await self.fec_client.get_candidate_totals(candidate_id, cycle=cycle)
            for total in totals:
                cycle_value = total.get('cycle') or total.get('two_year_transaction_period') or total.get('election_year') or 0
                data['financials'].append(FinancialSummary(
                    candidate_id=total.get('candidate_id', candidate_id),
                    cycle=cycle_value,
                    total_receipts=float(total.get('receipts', 0)),
                    total_disbursements=float(total.get('disbursements', 0)),
                    cash_on_hand=float(total.get('cash_on_hand_end_period', 0)),
                    total_contributions=float(total.get('contributions', 0)),
                    individual_contributions=float(total.get('individual_contributions', 0)),
                    pac_contributions=float(total.get('pac_contributions', 0)),
                    party_contributions=float(total.get('party_contributions', 0))
                ))
            
            # Get analysis data
            data['contribution_analysis'] = await self.analysis_service.analyze_contributions(
                candidate_id=candidate_id
            )
            data['expenditure_breakdown'] = await self.analysis_service.analyze_expenditures(
                candidate_id=candidate_id
            )
            data['employer_analysis'] = await self.analysis_service.analyze_by_employer(
                candidate_id=candidate_id
            )
            data['velocity'] = await self.analysis_service.analyze_velocity(
                candidate_id=candidate_id
            )
            data['fraud_analysis'] = await self.fraud_service.analyze_candidate(candidate_id)
            data['money_flow'] = await self.analysis_service.build_money_flow_graph(
                candidate_id=candidate_id,
                max_depth=2,
                min_amount=100.0
            )
        except Exception as e:
            # Continue with partial data
            pass
        
        return data
    
    async def collect_race_data(
        self,
        candidate_ids: List[str],
        office: str,
        state: str,
        district: Optional[str] = None,
        year: Optional[int] = None,
        cycle: Optional[int] = None
    ) -> Dict[str, Any]:
        """Collect data for multiple candidates in a race"""
        race_data = {
            'office': office,
            'state': state,
            'district': district,
            'year': year,
            'candidates': []
        }
        
        for candidate_id in candidate_ids:
            candidate_data = await self.collect_candidate_data(candidate_id, cycle=cycle)
            race_data['candidates'].append(candidate_data)
        
        return race_data
    
    def format_currency(self, amount: float) -> str:
        """Format currency for display"""
        if amount >= 1000000:
            return f"${amount/1000000:.2f}M"
        elif amount >= 1000:
            return f"${amount/1000:.2f}K"
        else:
            return f"${amount:.2f}"
    
    def generate_contribution_chart(self, analysis: ContributionAnalysis) -> BytesIO:
        """Generate contribution over time chart"""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        dates = sorted(analysis.contributions_by_date.keys())
        amounts = [analysis.contributions_by_date[d] for d in dates]
        
        ax.plot(dates, amounts, marker='o', linewidth=2, markersize=4)
        ax.set_title('Contributions Over Time', fontsize=14, fontweight='bold')
        ax.set_xlabel('Date', fontsize=12)
        ax.set_ylabel('Amount ($)', fontsize=12)
        ax.grid(True, alpha=0.3)
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        
        buf = BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf
    
    def generate_contribution_state_chart(self, analysis: ContributionAnalysis) -> BytesIO:
        """Generate contributions by state chart"""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        states = list(analysis.contributions_by_state.keys())[:10]  # Top 10
        amounts = [analysis.contributions_by_state[s] for s in states]
        
        ax.barh(states, amounts, color='steelblue')
        ax.set_title('Top 10 States by Contributions', fontsize=14, fontweight='bold')
        ax.set_xlabel('Amount ($)', fontsize=12)
        ax.set_ylabel('State', fontsize=12)
        ax.grid(True, alpha=0.3, axis='x')
        plt.tight_layout()
        
        buf = BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf
    
    def generate_top_donors_chart(self, analysis: ContributionAnalysis) -> BytesIO:
        """Generate top donors chart"""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        top_donors = sorted(analysis.top_donors, key=lambda x: x.get('total', 0), reverse=True)[:10]
        names = [d.get('name', 'Unknown')[:30] for d in top_donors]
        amounts = [d.get('total', 0) for d in top_donors]
        
        ax.barh(names, amounts, color='darkgreen')
        ax.set_title('Top 10 Donors', fontsize=14, fontweight='bold')
        ax.set_xlabel('Amount ($)', fontsize=12)
        ax.set_ylabel('Donor', fontsize=12)
        ax.grid(True, alpha=0.3, axis='x')
        plt.tight_layout()
        
        buf = BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf
    
    def generate_expenditure_category_chart(self, breakdown: ExpenditureBreakdown) -> BytesIO:
        """Generate expenditure by category pie chart"""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        categories = list(breakdown.expenditures_by_category.keys())
        amounts = [breakdown.expenditures_by_category[c] for c in categories]
        
        if amounts:
            ax.pie(amounts, labels=categories, autopct='%1.1f%%', startangle=90)
            ax.set_title('Expenditures by Category', fontsize=14, fontweight='bold')
        
        plt.tight_layout()
        
        buf = BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf
    
    def generate_expenditure_time_chart(self, breakdown: ExpenditureBreakdown) -> BytesIO:
        """Generate expenditures over time chart"""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        dates = sorted(breakdown.expenditures_by_date.keys())
        amounts = [breakdown.expenditures_by_date[d] for d in dates]
        
        ax.plot(dates, amounts, marker='o', linewidth=2, markersize=4, color='red')
        ax.set_title('Expenditures Over Time', fontsize=14, fontweight='bold')
        ax.set_xlabel('Date', fontsize=12)
        ax.set_ylabel('Amount ($)', fontsize=12)
        ax.grid(True, alpha=0.3)
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        
        buf = BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf
    
    def generate_fraud_radar_chart(self, fraud: FraudAnalysis) -> BytesIO:
        """Generate fraud pattern radar chart"""
        fig, ax = plt.subplots(figsize=(10, 6), subplot_kw=dict(projection='polar'))
        
        pattern_types = {}
        for pattern in fraud.patterns:
            ptype = pattern.pattern_type
            if ptype not in pattern_types:
                pattern_types[ptype] = {'count': 0, 'total': 0.0}
            pattern_types[ptype]['count'] += 1
            pattern_types[ptype]['total'] += pattern.total_amount
        
        if pattern_types:
            categories = list(pattern_types.keys())
            values = [pattern_types[c]['count'] for c in categories]
            
            angles = np.linspace(0, 2 * np.pi, len(categories), endpoint=False).tolist()
            values += values[:1]  # Complete the circle
            angles += angles[:1]
            
            ax.plot(angles, values, 'o-', linewidth=2)
            ax.fill(angles, values, alpha=0.25)
            ax.set_xticks(angles[:-1])
            ax.set_xticklabels(categories)
            ax.set_title('Fraud Pattern Analysis', fontsize=14, fontweight='bold', pad=20)
        
        plt.tight_layout()
        
        buf = BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf
    
    def generate_comparison_chart(
        self,
        comparison_data: List[Dict[str, Any]],
        metric: str,
        title: str
    ) -> BytesIO:
        """Generate comparison chart for race reports"""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        names = []
        values = []
        
        for comp in comparison_data:
            name = comp.get('name', 'Unknown')
            financials = comp.get('financials')
            if financials:
                names.append(name[:30])
                if metric == 'receipts':
                    values.append(financials.total_receipts)
                elif metric == 'cash':
                    values.append(financials.cash_on_hand)
                elif metric == 'disbursements':
                    values.append(financials.total_disbursements)
                elif metric == 'individual':
                    values.append(financials.individual_contributions)
                elif metric == 'pac':
                    values.append(financials.pac_contributions)
                else:
                    values.append(0)
        
        if values:
            ax.bar(names, values, color='steelblue')
            ax.set_title(title, fontsize=14, fontweight='bold')
            ax.set_ylabel('Amount ($)', fontsize=12)
            ax.set_xlabel('Candidate', fontsize=12)
            plt.xticks(rotation=45, ha='right')
            ax.grid(True, alpha=0.3, axis='y')
        
        plt.tight_layout()
        
        buf = BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf
    
    async def generate_pdf_report(self, data: Dict[str, Any], is_race: bool = False) -> BytesIO:
        """Generate PDF report"""
        buffer = BytesIO()
        
        if is_race:
            doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72,
                                   topMargin=72, bottomMargin=18)
        else:
            doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72,
                                   topMargin=72, bottomMargin=18)
        
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        if is_race:
            office_names = {'P': 'President', 'S': 'Senate', 'H': 'House of Representatives'}
            office_name = office_names.get(data['office'], data['office'])
            title_text = f"{office_name} Race - {data['state']}"
            if data.get('district'):
                title_text += f" District {data['district']}"
            if data.get('year'):
                title_text += f" ({data['year']})"
            story.append(Paragraph(title_text, title_style))
            story.append(Spacer(1, 12))
            
            # Candidates comparison table
            candidates = data.get('candidates', [])
            if candidates:
                story.append(Paragraph("Candidates Comparison", styles['Heading2']))
                story.append(Spacer(1, 12))
                
                # Get latest financials for each candidate
                comparison_data = []
                for cand_data in candidates:
                    candidate = cand_data.get('candidate')
                    financials = cand_data.get('financials', [])
                    if candidate and financials:
                        latest = max(financials, key=lambda x: x.cycle)
                        comparison_data.append({
                            'name': candidate.name,
                            'party': candidate.party or 'N/A',
                            'financials': latest
                        })
                
                if comparison_data:
                    table_data = [['Candidate', 'Party', 'Receipts', 'Disbursements', 'Cash on Hand', 'Individual', 'PAC']]
                    for comp in comparison_data:
                        fin = comp['financials']
                        table_data.append([
                            comp['name'][:30],
                            comp['party'],
                            self.format_currency(fin.total_receipts),
                            self.format_currency(fin.total_disbursements),
                            self.format_currency(fin.cash_on_hand),
                            self.format_currency(fin.individual_contributions),
                            self.format_currency(fin.pac_contributions)
                        ])
                    
                    table = Table(table_data, colWidths=[2*inch, 0.8*inch, 1*inch, 1*inch, 1*inch, 1*inch, 0.8*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTSIZE', (0, 1), (-1, -1), 8)
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 20))
                    
                    # Comparison charts
                    try:
                        chart_buf = self.generate_comparison_chart(comparison_data, 'receipts', 'Total Receipts Comparison')
                        img = Image(chart_buf, width=6*inch, height=3.6*inch)
                        story.append(img)
                        story.append(Spacer(1, 12))
                    except:
                        pass
        else:
            candidate = data.get('candidate')
            if candidate:
                title_text = candidate.name
                story.append(Paragraph(title_text, title_style))
                
                # Candidate info
                info_text = f"Office: {candidate.office or 'N/A'} | "
                info_text += f"Party: {candidate.party or 'N/A'} | "
                info_text += f"State: {candidate.state or 'N/A'}"
                if candidate.district:
                    info_text += f" | District: {candidate.district}"
                story.append(Paragraph(info_text, styles['Normal']))
                story.append(Spacer(1, 12))
        
        # Financial Summary
        financials = data.get('financials', [])
        if financials:
            story.append(Paragraph("Financial Summary", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            table_data = [['Cycle', 'Receipts', 'Disbursements', 'Cash on Hand', 'Individual', 'PAC']]
            for fin in sorted(financials, key=lambda x: x.cycle, reverse=True):
                table_data.append([
                    str(fin.cycle),
                    self.format_currency(fin.total_receipts),
                    self.format_currency(fin.total_disbursements),
                    self.format_currency(fin.cash_on_hand),
                    self.format_currency(fin.individual_contributions),
                    self.format_currency(fin.pac_contributions)
                ])
            
            table = Table(table_data, colWidths=[1*inch, 1.2*inch, 1.2*inch, 1.2*inch, 1.2*inch, 1*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(table)
            story.append(Spacer(1, 20))
        
        # Contribution Analysis
        contrib_analysis = data.get('contribution_analysis')
        if contrib_analysis and not is_race:
            story.append(Paragraph("Contribution Analysis", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            stats_text = f"Total Contributions: {self.format_currency(contrib_analysis.total_contributions)} | "
            stats_text += f"Total Contributors: {contrib_analysis.total_contributors} | "
            stats_text += f"Average Contribution: {self.format_currency(contrib_analysis.average_contribution)}"
            story.append(Paragraph(stats_text, styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Charts
            try:
                chart_buf = self.generate_contribution_chart(contrib_analysis)
                img = Image(chart_buf, width=6*inch, height=3.6*inch)
                story.append(img)
                story.append(Spacer(1, 12))
            except:
                pass
            
            # Top donors table
            if contrib_analysis.top_donors:
                story.append(Paragraph("Top Donors", styles['Heading3']))
                donor_data = [['Donor', 'Total', 'Count']]
                for donor in contrib_analysis.top_donors[:10]:
                    donor_data.append([
                        donor.get('name', 'Unknown')[:40],
                        self.format_currency(donor.get('total', 0)),
                        str(donor.get('count', 0))
                    ])
                donor_table = Table(donor_data, colWidths=[4*inch, 1.5*inch, 1*inch])
                donor_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(donor_table)
                story.append(Spacer(1, 20))
        
        # Expenditure Breakdown
        expenditure = data.get('expenditure_breakdown')
        if expenditure and not is_race:
            story.append(Paragraph("Expenditure Breakdown", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            stats_text = f"Total Expenditures: {self.format_currency(expenditure.total_expenditures)} | "
            stats_text += f"Total Transactions: {expenditure.total_transactions} | "
            stats_text += f"Average: {self.format_currency(expenditure.average_expenditure)}"
            story.append(Paragraph(stats_text, styles['Normal']))
            story.append(Spacer(1, 12))
            
            try:
                chart_buf = self.generate_expenditure_category_chart(expenditure)
                img = Image(chart_buf, width=6*inch, height=3.6*inch)
                story.append(img)
                story.append(Spacer(1, 12))
            except:
                pass
        
        # Fraud Analysis
        fraud = data.get('fraud_analysis')
        if fraud and fraud.patterns and not is_race:
            story.append(Paragraph("Fraud Analysis", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            risk_text = f"Risk Score: {fraud.risk_score:.2f} | "
            risk_text += f"Total Suspicious Amount: {self.format_currency(fraud.total_suspicious_amount)}"
            story.append(Paragraph(risk_text, styles['Normal']))
            story.append(Spacer(1, 12))
            
            try:
                chart_buf = self.generate_fraud_radar_chart(fraud)
                img = Image(chart_buf, width=6*inch, height=3.6*inch)
                story.append(img)
                story.append(Spacer(1, 12))
            except:
                pass
            
            # Fraud patterns table
            pattern_data = [['Pattern Type', 'Severity', 'Amount', 'Confidence']]
            for pattern in fraud.patterns[:10]:
                pattern_data.append([
                    pattern.pattern_type.replace('_', ' ').title(),
                    pattern.severity.upper(),
                    self.format_currency(pattern.total_amount),
                    f"{pattern.confidence_score:.2f}"
                ])
            pattern_table = Table(pattern_data, colWidths=[2.5*inch, 1*inch, 1.5*inch, 1*inch])
            pattern_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(pattern_table)
            story.append(Spacer(1, 20))
        
        # Footer
        story.append(Spacer(1, 20))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        footer_text = f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Data source: Federal Election Commission"
        story.append(Paragraph(footer_text, footer_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    async def generate_docx_report(self, data: Dict[str, Any], is_race: bool = False) -> BytesIO:
        """Generate DOCX report"""
        doc = Document()
        
        # Title
        if is_race:
            office_names = {'P': 'President', 'S': 'Senate', 'H': 'House of Representatives'}
            office_name = office_names.get(data['office'], data['office'])
            title_text = f"{office_name} Race - {data['state']}"
            if data.get('district'):
                title_text += f" District {data['district']}"
            if data.get('year'):
                title_text += f" ({data['year']})"
        else:
            candidate = data.get('candidate')
            title_text = candidate.name if candidate else "Candidate Report"
        
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if is_race:
            # Candidates comparison table
            candidates = data.get('candidates', [])
            if candidates:
                doc.add_heading('Candidates Comparison', 1)
                
                comparison_data = []
                for cand_data in candidates:
                    candidate = cand_data.get('candidate')
                    financials = cand_data.get('financials', [])
                    if candidate and financials:
                        latest = max(financials, key=lambda x: x.cycle)
                        comparison_data.append({
                            'name': candidate.name,
                            'party': candidate.party or 'N/A',
                            'financials': latest
                        })
                
                if comparison_data:
                    table = doc.add_table(rows=1, cols=7)
                    table.style = 'Light Grid Accent 1'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Candidate'
                    hdr_cells[1].text = 'Party'
                    hdr_cells[2].text = 'Receipts'
                    hdr_cells[3].text = 'Disbursements'
                    hdr_cells[4].text = 'Cash on Hand'
                    hdr_cells[5].text = 'Individual'
                    hdr_cells[6].text = 'PAC'
                    
                    for comp in comparison_data:
                        fin = comp['financials']
                        row_cells = table.add_row().cells
                        row_cells[0].text = comp['name'][:30]
                        row_cells[1].text = comp['party']
                        row_cells[2].text = self.format_currency(fin.total_receipts)
                        row_cells[3].text = self.format_currency(fin.total_disbursements)
                        row_cells[4].text = self.format_currency(fin.cash_on_hand)
                        row_cells[5].text = self.format_currency(fin.individual_contributions)
                        row_cells[6].text = self.format_currency(fin.pac_contributions)
                    
                    try:
                        chart_buf = self.generate_comparison_chart(comparison_data, 'receipts', 'Total Receipts Comparison')
                        doc.add_picture(chart_buf, width=Inches(6))
                    except:
                        pass
        else:
            candidate = data.get('candidate')
            if candidate:
                info_para = doc.add_paragraph()
                info_para.add_run(f"Office: {candidate.office or 'N/A'} | ")
                info_para.add_run(f"Party: {candidate.party or 'N/A'} | ")
                info_para.add_run(f"State: {candidate.state or 'N/A'}")
                if candidate.district:
                    info_para.add_run(f" | District: {candidate.district}")
        
        # Financial Summary
        financials = data.get('financials', [])
        if financials:
            doc.add_heading('Financial Summary', 1)
            
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Cycle'
            hdr_cells[1].text = 'Receipts'
            hdr_cells[2].text = 'Disbursements'
            hdr_cells[3].text = 'Cash on Hand'
            hdr_cells[4].text = 'Individual'
            hdr_cells[5].text = 'PAC'
            
            for fin in sorted(financials, key=lambda x: x.cycle, reverse=True):
                row_cells = table.add_row().cells
                row_cells[0].text = str(fin.cycle)
                row_cells[1].text = self.format_currency(fin.total_receipts)
                row_cells[2].text = self.format_currency(fin.total_disbursements)
                row_cells[3].text = self.format_currency(fin.cash_on_hand)
                row_cells[4].text = self.format_currency(fin.individual_contributions)
                row_cells[5].text = self.format_currency(fin.pac_contributions)
        
        # Contribution Analysis
        contrib_analysis = data.get('contribution_analysis')
        if contrib_analysis and not is_race:
            doc.add_heading('Contribution Analysis', 1)
            
            stats_para = doc.add_paragraph()
            stats_para.add_run(f"Total Contributions: {self.format_currency(contrib_analysis.total_contributions)} | ")
            stats_para.add_run(f"Total Contributors: {contrib_analysis.total_contributors} | ")
            stats_para.add_run(f"Average Contribution: {self.format_currency(contrib_analysis.average_contribution)}")
            
            # Add chart image
            try:
                chart_buf = self.generate_contribution_chart(contrib_analysis)
                doc.add_picture(chart_buf, width=Inches(6))
            except:
                pass
            
            # Top donors
            if contrib_analysis.top_donors:
                doc.add_heading('Top Donors', 2)
                donor_table = doc.add_table(rows=1, cols=3)
                donor_table.style = 'Light Grid Accent 1'
                hdr = donor_table.rows[0].cells
                hdr[0].text = 'Donor'
                hdr[1].text = 'Total'
                hdr[2].text = 'Count'
                
                for donor in contrib_analysis.top_donors[:10]:
                    row = donor_table.add_row().cells
                    row[0].text = donor.get('name', 'Unknown')[:40]
                    row[1].text = self.format_currency(donor.get('total', 0))
                    row[2].text = str(donor.get('count', 0))
        
        # Expenditure Breakdown
        expenditure = data.get('expenditure_breakdown')
        if expenditure and not is_race:
            doc.add_heading('Expenditure Breakdown', 1)
            
            stats_para = doc.add_paragraph()
            stats_para.add_run(f"Total Expenditures: {self.format_currency(expenditure.total_expenditures)} | ")
            stats_para.add_run(f"Total Transactions: {expenditure.total_transactions} | ")
            stats_para.add_run(f"Average: {self.format_currency(expenditure.average_expenditure)}")
            
            try:
                chart_buf = self.generate_expenditure_category_chart(expenditure)
                doc.add_picture(chart_buf, width=Inches(6))
            except:
                pass
        
        # Fraud Analysis
        fraud = data.get('fraud_analysis')
        if fraud and fraud.patterns and not is_race:
            doc.add_heading('Fraud Analysis', 1)
            
            risk_para = doc.add_paragraph()
            risk_para.add_run(f"Risk Score: {fraud.risk_score:.2f} | ")
            risk_para.add_run(f"Total Suspicious Amount: {self.format_currency(fraud.total_suspicious_amount)}")
            
            try:
                chart_buf = self.generate_fraud_radar_chart(fraud)
                doc.add_picture(chart_buf, width=Inches(6))
            except:
                pass
            
            # Fraud patterns table
            pattern_table = doc.add_table(rows=1, cols=4)
            pattern_table.style = 'Light Grid Accent 1'
            hdr = pattern_table.rows[0].cells
            hdr[0].text = 'Pattern Type'
            hdr[1].text = 'Severity'
            hdr[2].text = 'Amount'
            hdr[3].text = 'Confidence'
            
            for pattern in fraud.patterns[:10]:
                row = pattern_table.add_row().cells
                row[0].text = pattern.pattern_type.replace('_', ' ').title()
                row[1].text = pattern.severity.upper()
                row[2].text = self.format_currency(pattern.total_amount)
                row[3].text = f"{pattern.confidence_score:.2f}"
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(
            f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Data source: Federal Election Commission"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    async def generate_markdown_report(self, data: Dict[str, Any], is_race: bool = False) -> str:
        """Generate Markdown report"""
        md_lines = []
        
        # Title
        if is_race:
            office_names = {'P': 'President', 'S': 'Senate', 'H': 'House of Representatives'}
            office_name = office_names.get(data['office'], data['office'])
            title = f"# {office_name} Race - {data['state']}"
            if data.get('district'):
                title += f" District {data['district']}"
            if data.get('year'):
                title += f" ({data['year']})"
            md_lines.append(title)
            md_lines.append("")
            
            # Candidates comparison table
            candidates = data.get('candidates', [])
            if candidates:
                md_lines.append("## Candidates Comparison")
                md_lines.append("")
                
                comparison_data = []
                for cand_data in candidates:
                    candidate = cand_data.get('candidate')
                    financials = cand_data.get('financials', [])
                    if candidate and financials:
                        latest = max(financials, key=lambda x: x.cycle)
                        comparison_data.append({
                            'name': candidate.name,
                            'party': candidate.party or 'N/A',
                            'financials': latest
                        })
                
                if comparison_data:
                    md_lines.append("| Candidate | Party | Receipts | Disbursements | Cash on Hand | Individual | PAC |")
                    md_lines.append("|-----------|-------|----------|---------------|--------------|------------|-----|")
                    for comp in comparison_data:
                        fin = comp['financials']
                        md_lines.append(
                            f"| {comp['name'][:30]} | {comp['party']} | "
                            f"{self.format_currency(fin.total_receipts)} | "
                            f"{self.format_currency(fin.total_disbursements)} | "
                            f"{self.format_currency(fin.cash_on_hand)} | "
                            f"{self.format_currency(fin.individual_contributions)} | "
                            f"{self.format_currency(fin.pac_contributions)} |"
                        )
                    md_lines.append("")
        else:
            candidate = data.get('candidate')
            if candidate:
                md_lines.append(f"# {candidate.name}")
                md_lines.append("")
                info = f"**Office:** {candidate.office or 'N/A'} | "
                info += f"**Party:** {candidate.party or 'N/A'} | "
                info += f"**State:** {candidate.state or 'N/A'}"
                if candidate.district:
                    info += f" | **District:** {candidate.district}"
                md_lines.append(info)
                md_lines.append("")
        
        # Financial Summary
        financials = data.get('financials', [])
        if financials:
            md_lines.append("## Financial Summary")
            md_lines.append("")
            md_lines.append("| Cycle | Receipts | Disbursements | Cash on Hand | Individual | PAC |")
            md_lines.append("|-------|----------|---------------|--------------|------------|-----|")
            for fin in sorted(financials, key=lambda x: x.cycle, reverse=True):
                md_lines.append(
                    f"| {fin.cycle} | {self.format_currency(fin.total_receipts)} | "
                    f"{self.format_currency(fin.total_disbursements)} | {self.format_currency(fin.cash_on_hand)} | "
                    f"{self.format_currency(fin.individual_contributions)} | {self.format_currency(fin.pac_contributions)} |"
                )
            md_lines.append("")
        
        # Contribution Analysis
        contrib_analysis = data.get('contribution_analysis')
        if contrib_analysis and not is_race:
            md_lines.append("## Contribution Analysis")
            md_lines.append("")
            md_lines.append(
                f"- **Total Contributions:** {self.format_currency(contrib_analysis.total_contributions)}"
            )
            md_lines.append(f"- **Total Contributors:** {contrib_analysis.total_contributors}")
            md_lines.append(
                f"- **Average Contribution:** {self.format_currency(contrib_analysis.average_contribution)}"
            )
            md_lines.append("")
            
            # Top donors
            if contrib_analysis.top_donors:
                md_lines.append("### Top Donors")
                md_lines.append("")
                md_lines.append("| Donor | Total | Count |")
                md_lines.append("|-------|-------|-------|")
                for donor in contrib_analysis.top_donors[:10]:
                    md_lines.append(
                        f"| {donor.get('name', 'Unknown')[:40]} | "
                        f"{self.format_currency(donor.get('total', 0))} | {donor.get('count', 0)} |"
                    )
                md_lines.append("")
        
        # Expenditure Breakdown
        expenditure = data.get('expenditure_breakdown')
        if expenditure and not is_race:
            md_lines.append("## Expenditure Breakdown")
            md_lines.append("")
            md_lines.append(f"- **Total Expenditures:** {self.format_currency(expenditure.total_expenditures)}")
            md_lines.append(f"- **Total Transactions:** {expenditure.total_transactions}")
            md_lines.append(f"- **Average:** {self.format_currency(expenditure.average_expenditure)}")
            md_lines.append("")
        
        # Fraud Analysis
        fraud = data.get('fraud_analysis')
        if fraud and fraud.patterns and not is_race:
            md_lines.append("## Fraud Analysis")
            md_lines.append("")
            md_lines.append(f"- **Risk Score:** {fraud.risk_score:.2f}")
            md_lines.append(f"- **Total Suspicious Amount:** {self.format_currency(fraud.total_suspicious_amount)}")
            md_lines.append("")
            
            md_lines.append("### Fraud Patterns")
            md_lines.append("")
            md_lines.append("| Pattern Type | Severity | Amount | Confidence |")
            md_lines.append("|--------------|----------|--------|------------|")
            for pattern in fraud.patterns[:10]:
                md_lines.append(
                    f"| {pattern.pattern_type.replace('_', ' ').title()} | "
                    f"{pattern.severity.upper()} | {self.format_currency(pattern.total_amount)} | "
                    f"{pattern.confidence_score:.2f} |"
                )
            md_lines.append("")
        
        # Footer
        md_lines.append("---")
        md_lines.append(
            f"*Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
            f"Data source: Federal Election Commission*"
        )
        
        return "\n".join(md_lines)
    
    async def generate_csv_export(
        self,
        data: Dict[str, Any],
        is_race: bool = False
    ) -> BytesIO:
        """Generate CSV export of candidate or race data"""
        output = StringIO()
        writer = csv.writer(output)
        
        if is_race:
            # Race comparison CSV
            writer.writerow(['Race Analysis Report'])
            writer.writerow(['Office', data.get('office', 'N/A')])
            writer.writerow(['State', data.get('state', 'N/A')])
            if data.get('district'):
                writer.writerow(['District', data.get('district')])
            if data.get('year'):
                writer.writerow(['Year', data.get('year')])
            writer.writerow([])
            
            # Candidates comparison
            writer.writerow(['Candidates Comparison'])
            writer.writerow([
                'Candidate Name', 'Party', 'Cycle', 'Total Receipts',
                'Total Disbursements', 'Cash on Hand', 'Individual Contributions',
                'PAC Contributions', 'Party Contributions'
            ])
            
            for cand_data in data.get('candidates', []):
                candidate = cand_data.get('candidate')
                financials = cand_data.get('financials', [])
                if candidate and financials:
                    for fin in financials:
                        writer.writerow([
                            candidate.name,
                            candidate.party or 'N/A',
                            fin.cycle,
                            fin.total_receipts,
                            fin.total_disbursements,
                            fin.cash_on_hand,
                            fin.individual_contributions,
                            fin.pac_contributions,
                            fin.party_contributions
                        ])
        else:
            # Single candidate CSV
            candidate = data.get('candidate')
            if candidate:
                writer.writerow(['Candidate Information'])
                writer.writerow(['Name', candidate.name])
                writer.writerow(['Office', candidate.office or 'N/A'])
                writer.writerow(['Party', candidate.party or 'N/A'])
                writer.writerow(['State', candidate.state or 'N/A'])
                if candidate.district:
                    writer.writerow(['District', candidate.district])
                writer.writerow([])
            
            # Financial Summary
            financials = data.get('financials', [])
            if financials:
                writer.writerow(['Financial Summary'])
                writer.writerow([
                    'Cycle', 'Total Receipts', 'Total Disbursements',
                    'Cash on Hand', 'Total Contributions', 'Individual Contributions',
                    'PAC Contributions', 'Party Contributions'
                ])
                for fin in sorted(financials, key=lambda x: x.cycle, reverse=True):
                    writer.writerow([
                        fin.cycle,
                        fin.total_receipts,
                        fin.total_disbursements,
                        fin.cash_on_hand,
                        fin.total_contributions,
                        fin.individual_contributions,
                        fin.pac_contributions,
                        fin.party_contributions
                    ])
                writer.writerow([])
            
            # Contribution Analysis
            contrib_analysis = data.get('contribution_analysis')
            if contrib_analysis:
                writer.writerow(['Contribution Analysis'])
                writer.writerow(['Metric', 'Value'])
                writer.writerow(['Total Contributions', contrib_analysis.total_contributions])
                writer.writerow(['Total Contributors', contrib_analysis.total_contributors])
                writer.writerow(['Average Contribution', contrib_analysis.average_contribution])
                writer.writerow([])
                
                # Top Donors
                if contrib_analysis.top_donors:
                    writer.writerow(['Top Donors'])
                    writer.writerow(['Donor Name', 'Total Amount', 'Count'])
                    for donor in contrib_analysis.top_donors[:50]:
                        writer.writerow([
                            donor.get('name', 'Unknown'),
                            donor.get('total', 0),
                            donor.get('count', 0)
                        ])
                    writer.writerow([])
                
                # Contributions by State
                if contrib_analysis.contributions_by_state:
                    writer.writerow(['Contributions by State'])
                    writer.writerow(['State', 'Total Amount'])
                    for state, amount in sorted(
                        contrib_analysis.contributions_by_state.items(),
                        key=lambda x: x[1],
                        reverse=True
                    ):
                        writer.writerow([state, amount])
                    writer.writerow([])
        
        # Footer
        writer.writerow([])
        writer.writerow([f'Report generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'])
        writer.writerow(['Data source: Federal Election Commission'])
        
        output.seek(0)
        return BytesIO(output.getvalue().encode('utf-8'))
    
    async def generate_excel_export(
        self,
        data: Dict[str, Any],
        is_race: bool = False
    ) -> BytesIO:
        """Generate Excel export of candidate or race data"""
        wb = Workbook()
        ws = wb.active
        
        # Header style
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        header_alignment = Alignment(horizontal="center", vertical="center")
        
        # Title style
        title_font = Font(bold=True, size=14)
        
        row = 1
        
        if is_race:
            # Race comparison Excel
            ws.title = "Race Analysis"
            ws['A1'] = 'Race Analysis Report'
            ws['A1'].font = title_font
            row = 2
            
            ws[f'A{row}'] = 'Office'
            ws[f'B{row}'] = data.get('office', 'N/A')
            row += 1
            ws[f'A{row}'] = 'State'
            ws[f'B{row}'] = data.get('state', 'N/A')
            row += 1
            if data.get('district'):
                ws[f'A{row}'] = 'District'
                ws[f'B{row}'] = data.get('district')
                row += 1
            if data.get('year'):
                ws[f'A{row}'] = 'Year'
                ws[f'B{row}'] = data.get('year')
                row += 1
            row += 1
            
            # Candidates comparison
            ws[f'A{row}'] = 'Candidates Comparison'
            ws[f'A{row}'].font = title_font
            row += 1
            
            headers = [
                'Candidate Name', 'Party', 'Cycle', 'Total Receipts',
                'Total Disbursements', 'Cash on Hand', 'Individual Contributions',
                'PAC Contributions', 'Party Contributions'
            ]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col)
                cell.value = header
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = header_alignment
            row += 1
            
            for cand_data in data.get('candidates', []):
                candidate = cand_data.get('candidate')
                financials = cand_data.get('financials', [])
                if candidate and financials:
                    for fin in financials:
                        ws.cell(row=row, column=1, value=candidate.name)
                        ws.cell(row=row, column=2, value=candidate.party or 'N/A')
                        ws.cell(row=row, column=3, value=fin.cycle)
                        ws.cell(row=row, column=4, value=fin.total_receipts)
                        ws.cell(row=row, column=5, value=fin.total_disbursements)
                        ws.cell(row=row, column=6, value=fin.cash_on_hand)
                        ws.cell(row=row, column=7, value=fin.individual_contributions)
                        ws.cell(row=row, column=8, value=fin.pac_contributions)
                        ws.cell(row=row, column=9, value=fin.party_contributions)
                        row += 1
        else:
            # Single candidate Excel
            candidate = data.get('candidate')
            if candidate:
                ws.title = candidate.name[:31]  # Excel sheet name limit
                ws['A1'] = 'Candidate Information'
                ws['A1'].font = title_font
                row = 2
                
                info_data = [
                    ('Name', candidate.name),
                    ('Office', candidate.office or 'N/A'),
                    ('Party', candidate.party or 'N/A'),
                    ('State', candidate.state or 'N/A'),
                ]
                if candidate.district:
                    info_data.append(('District', candidate.district))
                
                for label, value in info_data:
                    ws.cell(row=row, column=1, value=label)
                    ws.cell(row=row, column=2, value=value)
                    row += 1
                row += 1
            
            # Financial Summary
            financials = data.get('financials', [])
            if financials:
                ws[f'A{row}'] = 'Financial Summary'
                ws[f'A{row}'].font = title_font
                row += 1
                
                headers = [
                    'Cycle', 'Total Receipts', 'Total Disbursements',
                    'Cash on Hand', 'Total Contributions', 'Individual Contributions',
                    'PAC Contributions', 'Party Contributions'
                ]
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=row, column=col)
                    cell.value = header
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = header_alignment
                row += 1
                
                for fin in sorted(financials, key=lambda x: x.cycle, reverse=True):
                    ws.cell(row=row, column=1, value=fin.cycle)
                    ws.cell(row=row, column=2, value=fin.total_receipts)
                    ws.cell(row=row, column=3, value=fin.total_disbursements)
                    ws.cell(row=row, column=4, value=fin.cash_on_hand)
                    ws.cell(row=row, column=5, value=fin.total_contributions)
                    ws.cell(row=row, column=6, value=fin.individual_contributions)
                    ws.cell(row=row, column=7, value=fin.pac_contributions)
                    ws.cell(row=row, column=8, value=fin.party_contributions)
                    row += 1
                row += 1
            
            # Contribution Analysis
            contrib_analysis = data.get('contribution_analysis')
            if contrib_analysis:
                ws[f'A{row}'] = 'Contribution Analysis'
                ws[f'A{row}'].font = title_font
                row += 1
                
                ws.cell(row=row, column=1, value='Metric')
                ws.cell(row=row, column=2, value='Value')
                for col in [1, 2]:
                    cell = ws.cell(row=row, column=col)
                    cell.fill = header_fill
                    cell.font = header_font
                row += 1
                
                ws.cell(row=row, column=1, value='Total Contributions')
                ws.cell(row=row, column=2, value=contrib_analysis.total_contributions)
                row += 1
                ws.cell(row=row, column=1, value='Total Contributors')
                ws.cell(row=row, column=2, value=contrib_analysis.total_contributors)
                row += 1
                ws.cell(row=row, column=1, value='Average Contribution')
                ws.cell(row=row, column=2, value=contrib_analysis.average_contribution)
                row += 2
                
                # Top Donors
                if contrib_analysis.top_donors:
                    ws[f'A{row}'] = 'Top Donors'
                    ws[f'A{row}'].font = title_font
                    row += 1
                    
                    headers = ['Donor Name', 'Total Amount', 'Count']
                    for col, header in enumerate(headers, 1):
                        cell = ws.cell(row=row, column=col)
                        cell.value = header
                        cell.fill = header_fill
                        cell.font = header_font
                        cell.alignment = header_alignment
                    row += 1
                    
                    for donor in contrib_analysis.top_donors[:50]:
                        ws.cell(row=row, column=1, value=donor.get('name', 'Unknown'))
                        ws.cell(row=row, column=2, value=donor.get('total', 0))
                        ws.cell(row=row, column=3, value=donor.get('count', 0))
                        row += 1
                    row += 1
                
                # Contributions by State
                if contrib_analysis.contributions_by_state:
                    ws[f'A{row}'] = 'Contributions by State'
                    ws[f'A{row}'].font = title_font
                    row += 1
                    
                    headers = ['State', 'Total Amount']
                    for col, header in enumerate(headers, 1):
                        cell = ws.cell(row=row, column=col)
                        cell.value = header
                        cell.fill = header_fill
                        cell.font = header_font
                        cell.alignment = header_alignment
                    row += 1
                    
                    for state, amount in sorted(
                        contrib_analysis.contributions_by_state.items(),
                        key=lambda x: x[1],
                        reverse=True
                    ):
                        ws.cell(row=row, column=1, value=state)
                        ws.cell(row=row, column=2, value=amount)
                        row += 1
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Footer
        row += 2
        ws.cell(row=row, column=1, value=f'Report generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        ws.cell(row=row + 1, column=1, value='Data source: Federal Election Commission')
        
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        return output

